"""Report Export Tool - Support Markdown, Word, PDF formats

Dependence on installation:
Pip install pypandoc markdown

The PDF export requires additional tools:
https://wkhtmltopdf.org/downloads.html
- Or LaTeX: https://www.latex-project.org/get/
"""

import logging
import os
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

#Check for dependency availability
try:
    import markdown
    import pypandoc

    #Check if pandoc is available
    try:
        pypandoc.get_pandoc_version()
        PANDOC_AVAILABLE = True
        logger.info("Pandoc")
    except OSError:
        PANDOC_AVAILABLE = False
        logger.warning("⚠️Pandoc is not available, Word and PDF export will not be available")

    EXPORT_AVAILABLE = True
except ImportError as e:
    EXPORT_AVAILABLE = False
    PANDOC_AVAILABLE = False
    logger.warning(f"Export dependency package is missing:{e}")
    logger.info("Please install: pip initial pypandoc markdown")

#Check pdfkit (sole PDF generation tool)
PDFKIT_AVAILABLE = False
PDFKIT_ERROR = None

try:
    import pdfkit
    #Check Wkhtmltopdf for installation
    try:
        pdfkit.configuration()
        PDFKIT_AVAILABLE = True
        logger.info("pdfkit + wkhtmltopdf is available (PDF Generation Tool)")
    except Exception as e:
        PDFKIT_ERROR = str(e)
        logger.warning("⚠️ wkhtmltopdf is not installed, PDF export is not available")
        logger.info("Installation methods: https://wkhtmltopdf.org/downloads.html")
except ImportError:
    logger.warning("pdfkit is not installed and PDF export is not available")
    logger.info("Installation: pip install pdfkit")
except Exception as e:
    PDFKIT_ERROR = str(e)
    logger.warning(f"pdfkit has failed:{e}")


class ReportExporter:
    """Report Exporter - Support Markdown, Word, PDF formats"""

    def __init__(self):
        self.export_available = EXPORT_AVAILABLE
        self.pandoc_available = PANDOC_AVAILABLE
        self.pdfkit_available = PDFKIT_AVAILABLE

        logger.info("Initialization of ReportExporter:")
        logger.info(f"  - export_available: {self.export_available}")
        logger.info(f"  - pandoc_available: {self.pandoc_available}")
        logger.info(f"  - pdfkit_available: {self.pdfkit_available}")
    
    def generate_markdown_report(self, report_doc: Dict[str, Any]) -> str:
        """Generate Markdown format reports"""
        logger.info("Make Markdown report...")
        
        stock_symbol = report_doc.get("stock_symbol", "unknown")
        analysis_date = report_doc.get("analysis_date", "")
        analysts = report_doc.get("analysts", [])
        research_depth = report_doc.get("research_depth", 1)
        reports = report_doc.get("reports", {})
        summary = report_doc.get("summary", "")
        
        content_parts = []
        
        #Title and meta-information
        content_parts.append(f"# {stock_symbol} 股票分析报告")
        content_parts.append("")
        content_parts.append(f"**分析日期**: {analysis_date}")
        if analysts:
            content_parts.append(f"**分析师**: {', '.join(analysts)}")
        content_parts.append(f"**研究深度**: {research_depth}")
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        
        #Executive summary
        if summary:
            content_parts.append("## 📊 执行摘要")
            content_parts.append("")
            content_parts.append(summary)
            content_parts.append("")
            content_parts.append("---")
            content_parts.append("")
        
        #Module content
        module_order = [
            "company_overview",
            "financial_analysis", 
            "technical_analysis",
            "market_analysis",
            "risk_analysis",
            "valuation_analysis",
            "investment_recommendation"
        ]
        
        module_titles = {
            "company_overview": "🏢 公司概况",
            "financial_analysis": "💰 财务分析",
            "technical_analysis": "📈 技术分析",
            "market_analysis": "🌍 市场分析",
            "risk_analysis": "⚠️ 风险分析",
            "valuation_analysis": "💎 估值分析",
            "investment_recommendation": "🎯 投资建议"
        }
        
        #Add modules sequentially
        for module_key in module_order:
            if module_key in reports:
                module_content = reports[module_key]
                if isinstance(module_content, str) and module_content.strip():
                    title = module_titles.get(module_key, module_key)
                    content_parts.append(f"## {title}")
                    content_parts.append("")
                    content_parts.append(module_content)
                    content_parts.append("")
                    content_parts.append("---")
                    content_parts.append("")
        
        #Add other unlisted modules
        for module_key, module_content in reports.items():
            if module_key not in module_order:
                if isinstance(module_content, str) and module_content.strip():
                    content_parts.append(f"## {module_key}")
                    content_parts.append("")
                    content_parts.append(module_content)
                    content_parts.append("")
                    content_parts.append("---")
                    content_parts.append("")
        
        #Footer
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        content_parts.append("*本报告由 TradingAgents-CN 自动生成*")
        content_parts.append("")
        
        markdown_content = "\n".join(content_parts)
        logger.info(f"Markdown report generated, length:{len(markdown_content)}Character")
        
        return markdown_content
    
    def _clean_markdown_for_pandoc(self, md_content: str) -> str:
        """Clear Markdown content to avoid pandoc parsing problems"""
        import re

        #Remove content that could cause YAML resolution problems
        #If there's "--" at the beginning, add an empty line to the front
        if md_content.strip().startswith("---"):
            md_content = "\n" + md_content

        #Remove possibly vertical HTML labels and styles
        #Remove wringing-mode-related styles
        md_content = re.sub(r'<[^>]*writing-mode[^>]*>', '', md_content, flags=re.IGNORECASE)
        md_content = re.sub(r'<[^>]*text-orientation[^>]*>', '', md_content, flags=re.IGNORECASE)

        #Remove <div> label 's syle properties (possibly containing vertical styles)
        md_content = re.sub(r'<div\s+style="[^"]*">', '<div>', md_content, flags=re.IGNORECASE)
        md_content = re.sub(r'<span\s+style="[^"]*">', '<span>', md_content, flags=re.IGNORECASE)

        #Remove potentially problematic HTML labels
        #Keep basic Markdown format and remove complex HTML
        md_content = re.sub(r'<style[^>]*>.*?</style>', '', md_content, flags=re.DOTALL | re.IGNORECASE)

        #Make sure all the paragraphs are in normal rows.
        #Add a clear line break around each paragraph to avoid Pandoc error
        lines = md_content.split('\n')
        cleaned_lines = []
        for line in lines:
            #Skip empty lines
            if not line.strip():
                cleaned_lines.append(line)
                continue

            #Markdown syntax if title, list, table, etc., maintain as is
            if line.strip().startswith(('#', '-', '*', '|', '>', '```', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                cleaned_lines.append(line)
            else:
                #Normal Paragraph: Ensure that no special characters cause vertical rows
                cleaned_lines.append(line)

        md_content = '\n'.join(cleaned_lines)

        return md_content

    def _create_pdf_css(self) -> str:
        """Create a PDF stylesheet to control the page break and text orientation of tables"""
        return """
<style>
/* 🔥 强制所有文本横排显示（修复中文竖排问题） */
* {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

body {
    writing-mode: horizontal-tb !important;
    direction: ltr !important;
}

/* 段落和文本 */
p, div, span, td, th, li {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

/* 表格样式 - 允许跨页 */
table {
    width: 100%;
    border-collapse: collapse;
    page-break-inside: auto;
    writing-mode: horizontal-tb !important;
}

/* 表格行 - 避免在行中间分页 */
tr {
    page-break-inside: avoid;
    page-break-after: auto;
}

/* 表头 - 在每页重复显示 */
thead {
    display: table-header-group;
}

/* 表格单元格 */
td, th {
    padding: 8px;
    border: 1px solid #ddd;
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

/* 表头样式 */
th {
    background-color: #f2f2f2;
    font-weight: bold;
}

/* 避免标题后立即分页 */
h1, h2, h3, h4, h5, h6 {
    page-break-after: avoid;
    writing-mode: horizontal-tb !important;
}

/* 避免在列表项中间分页 */
li {
    page-break-inside: avoid;
}

/* 代码块 */
pre, code {
    writing-mode: horizontal-tb !important;
    white-space: pre-wrap;
    word-wrap: break-word;
}
</style>
"""
    
    def generate_docx_report(self, report_doc: Dict[str, Any]) -> bytes:
        """Generate Word Document Format Report"""
        logger.info("Start generating Word documents...")

        if not self.pandoc_available:
            raise Exception("Pandoc 不可用，无法生成 Word 文档。请安装 pandoc 或使用 Markdown 格式导出。")

        #Generate Markdown content
        md_content = self.generate_markdown_report(report_doc)

        try:
            #Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                output_file = tmp_file.name

            logger.info(f"Temporary file path:{output_file}")

            #Pandoc Arguments
            extra_args = [
                '--from=markdown-yaml_metadata_block',  #Disable YAML metadata block resolution
                '--standalone',  #Generate a separate document
                '--wrap=preserve',  #Keep Line Break
                '--columns=120',  #Set column width
                '-M', 'lang=zh-CN',  #🔥 Specifying language to Chinese
                '-M', 'dir=ltr',  #🔥 Specify text direction from left to right
            ]

            #Clear Contents
            cleaned_content = self._clean_markdown_for_pandoc(md_content)

            #Convert to Word
            pypandoc.convert_text(
                cleaned_content,
                'docx',
                format='markdown',
                outputfile=output_file,
                extra_args=extra_args
            )

            logger.info("Pypandoc conversion complete")

            #🔥 Postprocessing: fixing text orientation in Word documents
            try:
                from docx import Document
                doc = Document(output_file)

                #Fix text orientation for all paragraphs
                for paragraph in doc.paragraphs:
                    #Set the paragraphs from left to right
                    if paragraph._element.pPr is not None:
                        #Remove possible vertical settings
                        for child in list(paragraph._element.pPr):
                            if 'textDirection' in child.tag or 'bidi' in child.tag:
                                paragraph._element.pPr.remove(child)

                #Fix Text Directions in Tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph._element.pPr is not None:
                                    for child in list(paragraph._element.pPr):
                                        if 'textDirection' in child.tag or 'bidi' in child.tag:
                                            paragraph._element.pPr.remove(child)

                #Save recovered documents
                doc.save(output_file)
                logger.info("Word Document Text Restoration Complete")
            except ImportError:
                logger.warning("⚠️ python-docx uninstalled, skipping text direction repair")
            except Exception as e:
                logger.warning(f"Could not close temporary folder: %s{e}")

            #Read generated files
            with open(output_file, 'rb') as f:
                docx_content = f.read()

            logger.info(f"Word document generation success, size:{len(docx_content)}Bytes")

            #Clear temporary files
            os.unlink(output_file)

            return docx_content
            
        except Exception as e:
            logger.error(f"Could not close temporary folder: %s{e}", exc_info=True)
            #Clear temporary files
            try:
                if 'output_file' in locals() and os.path.exists(output_file):
                    os.unlink(output_file)
            except:
                pass
            raise Exception(f"生成 Word 文档失败: {e}")
    
    def _markdown_to_html(self, md_content: str) -> str:
        """Convert Markdown to HTML"""
        import markdown

        #Configure Markdown Extensions
        extensions = [
            'markdown.extensions.tables',  #Table support
            'markdown.extensions.fenced_code',  #Block Support
            'markdown.extensions.nl2br',  #Line Break Support
        ]

        #Convert to HTML
        html_content = markdown.markdown(md_content, extensions=extensions)

        #Add HTML templates and styles
        #WeasyPrint Optimized CSS (delegate unsupported properties)
        html_template = f"""
<!DOCTYPE html>
<html lang="zh-CN" dir="ltr">
<head>
    <meta charset="UTF-8">
    <title>分析报告</title>
    <style>
        /* 基础样式 - 确保文本方向正确 */
        html {{
            direction: ltr;
        }}

        body {{
            font-family: "Noto Sans CJK SC", "Microsoft YaHei", "SimHei", "Arial", sans-serif;
            line-height: 1.8;
            color: #333;
            margin: 20mm;
            padding: 0;
            background: white;
            direction: ltr;
        }}

        /* 标题样式 */
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 1.5em;
            margin-bottom: 0.8em;
            font-weight: 600;
            page-break-after: avoid;
            direction: ltr;
        }}

        h1 {{
            font-size: 2em;
            border-bottom: 3px solid #3498db;
            padding-bottom: 0.3em;
            page-break-before: always;
        }}

        h1:first-child {{
            page-break-before: avoid;
        }}

        h2 {{
            font-size: 1.6em;
            border-bottom: 2px solid #bdc3c7;
            padding-bottom: 0.25em;
        }}

        h3 {{
            font-size: 1.3em;
            color: #34495e;
        }}

        /* 段落样式 */
        p {{
            margin: 0.8em 0;
            text-align: left;
            direction: ltr;
        }}

        /* 表格样式 - 优化分页 */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1.5em 0;
            font-size: 0.9em;
            direction: ltr;
        }}

        /* 表头在每页重复 */
        thead {{
            display: table-header-group;
        }}

        tbody {{
            display: table-row-group;
        }}

        /* 表格行避免跨页断开 */
        tr {{
            page-break-inside: avoid;
        }}

        th, td {{
            border: 1px solid #ddd;
            padding: 10px 12px;
            text-align: left;
            direction: ltr;
        }}

        th {{
            background-color: #3498db;
            color: white;
            font-weight: bold;
        }}

        tbody tr:nth-child(even) {{
            background-color: #f8f9fa;
        }}

        tbody tr:hover {{
            background-color: #e9ecef;
        }}

        /* 代码块样式 */
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: "Consolas", "Monaco", "Courier New", monospace;
            font-size: 0.9em;
            direction: ltr;
        }}

        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            border-left: 4px solid #3498db;
            page-break-inside: avoid;
            direction: ltr;
        }}

        pre code {{
            background-color: transparent;
            padding: 0;
        }}

        /* 列表样式 */
        ul, ol {{
            margin: 0.8em 0;
            padding-left: 2em;
            direction: ltr;
        }}

        li {{
            margin: 0.4em 0;
            direction: ltr;
        }}

        /* 强调文本 */
        strong, b {{
            font-weight: 700;
            color: #2c3e50;
        }}

        em, i {{
            font-style: italic;
            color: #555;
        }}

        /* 水平线 */
        hr {{
            border: none;
            border-top: 2px solid #ecf0f1;
            margin: 2em 0;
        }}

        /* 链接样式 */
        a {{
            color: #3498db;
            text-decoration: none;
        }}

        a:hover {{
            text-decoration: underline;
        }}

        /* 分页控制 */
        @page {{
            size: A4;
            margin: 20mm;

            @top-center {{
                content: "分析报告";
                font-size: 10pt;
                color: #999;
            }}

            @bottom-right {{
                content: "第 " counter(page) " 页";
                font-size: 10pt;
                color: #999;
            }}
        }}

        /* 避免孤行和寡行 */
        p, li {{
            orphans: 3;
            widows: 3;
        }}

        /* 图片样式 */
        img {{
            max-width: 100%;
            height: auto;
            page-break-inside: avoid;
        }}

        /* 引用块样式 */
        blockquote {{
            margin: 1em 0;
            padding: 0.5em 1em;
            border-left: 4px solid #3498db;
            background-color: #f8f9fa;
            font-style: italic;
            page-break-inside: avoid;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
        return html_template

    def _generate_pdf_with_pdfkit(self, html_content: str) -> bytes:
        """Generate PDF using pdfkit"""
        import pdfkit

        logger.info("Create PDF with pdfkit + wkhtmltopdf...")

        #Configure Options
        options = {
            'encoding': 'UTF-8',
            'enable-local-file-access': None,
            'page-size': 'A4',
            'margin-top': '20mm',
            'margin-right': '20mm',
            'margin-bottom': '20mm',
            'margin-left': '20mm',
        }

        #Generate PDF
        pdf_bytes = pdfkit.from_string(html_content, False, options=options)

        logger.info(f"pdfkit PDF generation success, size:{len(pdf_bytes)}Bytes")
        return pdf_bytes

    def generate_pdf_report(self, report_doc: Dict[str, Any]) -> bytes:
        """Generate PDF format reports (using pdfkit + wkhtmltopdf)"""
        logger.info("Start generating PDF documents...")

        #Check if pdfkit is available
        if not self.pdfkit_available:
            error_msg = (
                "pdfkit 不可用，无法生成 PDF。\n\n"
                "安装方法:\n"
                "1. 安装 pdfkit: pip install pdfkit\n"
                "2. 安装 wkhtmltopdf: https://wkhtmltopdf.org/downloads.html\n"
            )
            if PDFKIT_ERROR:
                error_msg += f"\n错误详情: {PDFKIT_ERROR}"

            logger.error(f"❌ {error_msg}")
            raise Exception(error_msg)

        #Generate Markdown content
        md_content = self.generate_markdown_report(report_doc)

        #Generate PDF using pdfkit
        try:
            html_content = self._markdown_to_html(md_content)
            return self._generate_pdf_with_pdfkit(html_content)
        except Exception as e:
            error_msg = f"PDF 生成失败: {e}"
            logger.error(f"❌ {error_msg}")
            raise Exception(error_msg)


#Create global export instance
report_exporter = ReportExporter()

